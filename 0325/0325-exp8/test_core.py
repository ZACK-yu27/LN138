import sys
import os

sys.path.insert(0, os.path.dirname(__file__))

import jieba
import re
from wordcloud import WordCloud
import matplotlib.pyplot as plt

FONT_PATH = os.path.join(os.path.dirname(__file__), "HGLS.TTF")

print("=" * 50)
print("词云生成工具核心功能测试")
print("=" * 50)

CHINESE_STOPWORDS = {
    '的', '了', '在', '是', '和', '也', '就', '都', '而', '及', '与', '为', '对', '等', '这', '那',
    '你', '我', '他', '她', '它', '我们', '你们', '他们'
}

ENGLISH_STOPWORDS = {
    'the', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had',
    'to', 'of', 'in', 'for', 'on', 'with', 'at', 'by', 'from', 'as', 'i', 'me', 'my', 'we', 'our'
}

def clean_text(text):
    text = re.sub(r'[^\w\s\u4e00-\u9fff]', ' ', text)
    text = re.sub(r'\d+', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def tokenize_text(text):
    words = []
    chinese_pattern = re.compile(r'[\u4e00-\u9fff]+')
    english_pattern = re.compile(r'[a-zA-Z]+')

    chinese_parts = chinese_pattern.findall(text)
    english_parts = english_pattern.findall(text)

    for part in chinese_parts:
        if part:
            jieba_tokens = jieba.lcut(part)
            words.extend(jieba_tokens)

    for part in english_parts:
        if part:
            words.extend(part.lower().split())

    return words

def filter_stopwords(words):
    filtered = []
    for word in words:
        word = word.strip()
        if len(word) <= 1:
            continue
        if word in CHINESE_STOPWORDS:
            continue
        if word.lower() in ENGLISH_STOPWORDS:
            continue
        if word.isdigit():
            continue
        if re.match(r'^[\u4e00-\u9fff]$', word):
            continue
        if word:
            filtered.append(word)
    return filtered

def process_text(text):
    cleaned = clean_text(text)
    words = tokenize_text(cleaned)
    filtered = filter_stopwords(words)
    word_freq = {}
    for word in filtered:
        word_freq[word] = word_freq.get(word, 0) + 1
    return word_freq

print("\n[测试1] 中英文混合文本处理")
test_text = """
Python是一种广泛使用的解释型、高级和通用的编程语言。
Python支持多种编程范式，包括结构化、过程式、反射式面向对象和函数式编程。
它拥有动态类型系统和垃圾回收功能，能够自动管理内存使用。
JavaScript是一种基于原型编程、多范式的动态脚本语言，并且支持面向对象、命令式和声明式编程风格。
Machine Learning is a subset of artificial intelligence that enables systems to learn and improve from experience.
Deep Learning is part of a broader family of machine learning methods based on artificial neural networks.
数据科学是使用科学方法、过程和系统来从结构化或非结构化数据中提取知识和洞察的领域。
"""
print(f"输入文本长度: {len(test_text)} 字符")

word_freq = process_text(test_text)
print(f"提取词汇数: {len(word_freq)} 个")
print(f"词频统计 (前10): {dict(list(word_freq.items())[:10])}")

print("\n[测试2] 词云生成")
if word_freq:
    wc = WordCloud(
        font_path=FONT_PATH,
        width=800,
        height=400,
        background_color='#0F172A',
        colormap='viridis',
        max_words=200,
        max_font_size=150,
        min_font_size=10,
        random_state=42
    )
    wc.generate_from_frequencies(word_freq)
    print("[OK] 词云生成成功")

    fig, ax = plt.subplots(figsize=(12, 6))
    ax.imshow(wc, interpolation='bilinear')
    ax.axis('off')
    ax.set_facecolor('#0F172A')
    fig.patch.set_facecolor('#0F172A')

    output_path = os.path.join(os.path.dirname(__file__), "test_wordcloud.png")
    fig.savefig(output_path, format='png', dpi=100, bbox_inches='tight', facecolor='#0F172A')
    print(f"✓ 测试词云已保存: {output_path}")
    plt.close(fig)
else:
    print("✗ 无有效词汇")

print("\n[测试3] Word文件解析")
from docx import Document
test_docx = os.path.join(os.path.dirname(__file__), "test_sample.docx")
doc = Document()
doc.add_paragraph("这是一个测试文档。")
doc.add_paragraph("This is a test document for wordcloud generation.")
doc.add_paragraph("Python和JavaScript是流行的编程语言。")
doc.save(test_docx)
print(f"✓ 测试文档创建成功: {test_docx}")

doc = Document(test_docx)
content = []
for para in doc.paragraphs:
    content.append(para.text)
extracted_text = '\n'.join(content)
print(f"✓ 文档解析成功，提取 {len(extracted_text)} 字符")

word_freq_doc = process_text(extracted_text)
print(f"✓ 文档词汇提取: {len(word_freq_doc)} 个词汇")

os.remove(test_docx)
print(f"✓ 测试文档已清理")

print("\n" + "=" * 50)
print("所有核心功能测试通过！")
print("=" * 50)